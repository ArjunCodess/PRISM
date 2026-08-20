"""Fill CJSJ official templates with the research cut. Layout comes from the templates."""

from __future__ import annotations

import re
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Emu
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
CJSJ = ROOT / "submissions" / "cjsj"
FIGS = ROOT / "docs" / "figures"
DOCX_TMPL = CJSJ / "CJSJ+Original+Research+Template+(1).docx"
PPTX_TMPL = CJSJ / "CJSJ+Figures+Template.pptx"
OUT_DOCX = CJSJ / "PrakashArjun_paper.docx"
OUT_PPTX = CJSJ / "PrakashArjun_figures.pptx"
UNPACK = CJSJ / "_unpack" / "pptx"

TITLE = (
    "Predictive information in conjunction data-message histories "
    "decays as time to closest approach decreases"
)

CAPTIONS = [
    (
        "horizon-decay.png",
        "Figure 1. Horizon MAE: extra information is largest at 72 h and nearly gone at 12 h.",
    ),
    (
        "error-anatomy.png",
        "Figure 2. Later movement versus residual. Spike at 0: no movement. Left tail: floor collapse.",
    ),
    (
        "coverage-calibration.png",
        "Figure 3. Nominal versus empirical coverage for bootstrap spread and split conformal.",
    ),
    (
        "dilution-probe.png",
        "Figure 4. Floor rate and mean |y \u2212 risk| by max-risk-gap quartile. Large gaps track the floor.",
    ),
]
ADD_SLIDE = Path(r"C:\Users\Arjun\.agents\skills\pptx\scripts\add_slide.py")
CLEAN_SLIDE = Path(r"C:\Users\Arjun\.agents\skills\pptx\scripts\clean.py")
PLACEHOLDER_CAPTION = "Figure 1.\tExample of a figure caption."


def _set_run_font(run, size=10, bold=False, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _p(doc, text, *, size=10, bold=False, italic=False, style=None, align=None, first=False, space_after=6):
    p = doc.add_paragraph()
    if style:
        try:
            p.style = style
        except KeyError:
            pass
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 0.95
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if first:
        p.paragraph_format.first_line_indent = Inches(0.2)
    else:
        p.paragraph_format.first_line_indent = Inches(0)
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def _heading(doc, text):
    return _p(doc, text, size=10, bold=True, style="Heading 1", space_after=4)


def _sub(doc, text):
    return _p(doc, text, size=10, bold=True, italic=True, space_after=2)


def _body(doc, text, first=True):
    return _p(doc, text, size=10, first=first, space_after=6)


def _caption(doc, text):
    return _p(doc, text, size=8, italic=True, space_after=8)


def _keep_with_next(paragraph) -> None:
    paragraph._p.get_or_add_pPr().append(OxmlElement("w:keepNext"))


def _table(doc, headers, rows):
    if doc.paragraphs:
        _keep_with_next(doc.paragraphs[-1])
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        _set_run_font(run, size=8, bold=True)
    hdr_trPr = table.rows[0]._tr.get_or_add_trPr()
    hdr_trPr.append(OxmlElement("w:tblHeader"))
    hdr_trPr.append(OxmlElement("w:cantSplit"))
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            _set_run_font(run, size=8)
        table.rows[ri + 1]._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def _figure(doc, filename, cap):
    path = FIGS / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(3.15))
    _caption(doc, cap)


def fill_docx() -> None:
    shutil.copy2(DOCX_TMPL, OUT_DOCX)
    doc = Document(str(OUT_DOCX))
    body = doc.element.body
    sect = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is not sect:
            body.remove(child)

    _p(doc, TITLE, size=14, bold=True, style="Title", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    _p(doc, "Arjun Vijay Prakash", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _p(
        doc,
        "City Montessori School, Kanpur Road Campus, Lucknow, India",
        size=10,
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )

    _heading(doc, "Abstract")
    _body(
        doc,
        "Conjunction Data Messages recompute collision probability as later tracking shrinks orbital uncertainty. An earlier forecast is useful only if the history available then contains information beyond copying the latest report. This study freezes the clock at 72, 48, 24, and 12 hours before closest approach on the public ESA Collision Avoidance Challenge archive under a leakage-safe split. Extra information is real, but it is a floor-collapse statistic and a long-horizon statistic, not a typical-event accuracy win and not a high-risk warning win. On 1659 held-out events, 1352 later reports sit at the dataset floor of −30, and persistence already has median absolute error 0. Unguarded XGBoost lowers mean absolute error from 5.080 to 2.809, yet floor-excluded MAE rises from 4.073 to 7.562 and a Wilcoxon test on paired absolute error does not reject equality (p=0.91). The learned advantage is 4.53 log units at 72 hours and 0.06 at 12 hours. Validation therefore selects a two-part floor hurdle (test MAE 2.109, F2=0). On ESA's 2167 official-test events, clipped persistence matches Uriot last-risk-prediction loss L=0.694; the floor hurdle lowers MAE but raises L to 70.5. Combined covariance volume tracks later movement; the max-risk gap tracks floor membership. This is a research prototype, not flight software.",
        first=False,
    )

    _heading(doc, "Introduction")
    _body(
        doc,
        "Two orbiting objects can pass close enough to generate a conjunction alert. The reported collision probability Pc is not a physical constant. It is recomputed as radar and optical observations shrink, or fail to shrink, the combined covariance [1], [2]. Waiting usually improves the estimate and leaves less time to plan. ESA posed the forecasting form of this problem in the 2019 Collision Avoidance Challenge: predict the final reported log10 Pc from messages with time to closest approach of at least two days [3], [4]. Persistence—copying the latest pre-cutoff risk—was a strong baseline. Only 12 of 96 teams beat it on ESA-style loss.",
    )
    _body(
        doc,
        "That competition does not say how any extra information depends on remaining time, or whether a mean-error win is a typical-event win. Physics-based Pc forecasting via expected covariance reduction [5] and operational work on probability dilution [6], [2] suggest a concrete picture: large early covariances can make Pc look safely small, then later updates either collapse risk to negligible or concentrate it. The target here is the later reported log10 Pc itself, under a frozen information cutoff, on public CDMs.",
    )
    _body(
        doc,
        "Four working hypotheses are measured. H1: extra predictive information beyond copying the latest report is largest at long horizons and approaches zero near closest approach. H2: most of the mean-error win is rare floor jumps, not typical events becoming more accurate. H3: extra information does not automatically improve high-risk decisions under the ESA class log10 Pc ≥ −6. H4: large combined covariance at T−48 predicts later |Δrisk|; the snapshot gap max_risk_estimate − risk predicts floor membership, not larger later movement. Mean error and ESA-style loss disagree on the validation-selected floor hurdle. That disagreement is part of the result.",
    )

    _heading(doc, "Methods")
    _sub(doc, "Participants/Organisms")
    _body(
        doc,
        "There are no human participants. Labeled CDMs from the ESA training archive comprise 162,634 rows and 13,154 events, 2015–2019, anonymized ESA-supported missions [4]. An event needs a message at or before the cutoff and a later labeled update (8,293 eligible events at 48 hours). This is not live ISRO, NASA, or operational catalogue data. Manoeuvre advice is out of scope.",
        first=False,
    )
    _sub(doc, "Experimental Design")
    _body(
        doc,
        "Train, validation, calibration, and test are event-disjoint (3,731 / 1,659 / 1,244 / 1,659). Seed 42 is the reported local split. Five grouped redraws use seeds 42–46 with no test-tuned knobs. Features use only rows with time to closest approach at or above the cutoff. The later log10 Pc is the label, never an input. Official-test inputs are the public Kelvins test file. Labels are the released final-risk column from Zenodo record 4463683 (25 January 2021), joined after freeze [4]. No hyperparameter was changed after seeing official-test scores.",
        first=False,
    )
    _body(
        doc,
        "The selected policy, chosen on validation by MAE, is a two-part floor hurdle: a classifier for P(y=−30) and a residual regressor fit only on non-floor training events (threshold 0.15; no −6 persistence guard). Split-conformal intervals are fit on calibration residuals of that point [7]. The model abstains when the 90% conformal band crosses −6 or a critical field is missing. Baselines are persistence, unguarded XGBoost, residual reconstruction, and a bootstrap median with a −6 persist guard.",
        first=False,
    )
    _sub(doc, "Data Analysis")
    _body(
        doc,
        "Scores are mean and median absolute error in log10 Pc and ESA-style loss L = MSE_HR / F2 on the −6 class [3]. Floor-excluded MAE is MAE off the −30 floor. A 95% interval on MAE(persistence) − MAE(model) uses 1,000 event-resamples. A two-sided Wilcoxon signed-rank test compares per-event absolute errors. Official-test L is the secondary evaluation (2,167 events, 150 high-risk). All numbers are from the frozen evaluation in metrics.json.",
        first=False,
    )

    _heading(doc, "Results/Discussion")
    _body(
        doc,
        "Table I reports the frozen local test. Of 1,659 events, 1,352 later reports sit at −30. Persistence median AE is 0. Residual XGBoost reaches test MAE 2.760, close to unguarded XGBoost (2.809). Both beat persistence on the mean. Neither beats it on the typical event: floor-excluded MAE rises, and Wilcoxon tests do not reject equality (p=0.91 and 0.92). That is H2. The floor hurdle is the validation winner (test MAE 2.109, median AE 0). Floor-excluded MAE rises to 9.311 versus persistence 4.073. F2=0 for the floor hurdle. ESA-style loss still ties at 0.167 for persistence and the guarded ensemble.",
        first=False,
    )
    _table(
        doc,
        ["System", "MAE", "Med.", "Non-fl.", "p"],
        [
            ["Persistence", "5.080", "0.000", "4.073", "—"],
            ["Unguarded XGB", "2.809", "0.554", "7.562", "0.91"],
            ["Residual XGB", "2.760", "0.453", "7.375", "0.92"],
            ["Floor hurdle", "2.109", "0.000", "9.311", "<10^−33"],
            ["Guarded ens.", "3.059", "0.473", "7.332", "0.31"],
        ],
    )
    _caption(doc, "Table I. Frozen local test at 48 h (n=1,659; 1,352 floor). MAE in log10 Pc. p is two-sided Wilcoxon on paired absolute error.")
    _figure(doc, *CAPTIONS[1])
    _body(
        doc,
        "Table II and Figure 1 support H1. The MAE advantage is 4.534 at 72 hours, 2.272 at 48 hours, 0.524 at 24 hours, and 0.060 at 12 hours. Across five grouped redraws, unguarded XGBoost MAE advantage is 2.34 ± 0.04. Leave-one-out on 66 local high-risk events: persistence is closer than residual reconstruction on 66/66.",
    )
    _table(
        doc,
        ["Horizon", "XGBoost", "Persist.", "Adv."],
        [
            ["72 h", "3.214", "7.748", "4.534"],
            ["48 h", "2.808", "5.080", "2.272"],
            ["24 h", "2.110", "2.634", "0.524"],
            ["12 h", "1.384", "1.444", "0.060"],
        ],
    )
    _caption(doc, "Table II. Single XGBoost versus persistence by horizon. MAE in log10 Pc.")
    _figure(doc, *CAPTIONS[0])
    _body(
        doc,
        "Around the selected floor hurdle the 90% conformal band covers 90.1% of the test (mean width 21.03). Validation chose conformal crossing −6 as the abstention rule (test coverage 90.4%; false reassurance 2 of 9). Bootstrap 90% bands around the guarded ensemble cover 47.7% of outcomes (Figure 3); they are model spread.",
    )
    _figure(doc, *CAPTIONS[2])
    _body(
        doc,
        "Table III is the one-shot official-test score (150 positives). Persistence L=0.694 matches Uriot last-risk-prediction [3]. Residual and floor models reduce MAE (3.476 and 3.177 versus 5.209) but F2 collapses and L rises to 104 and 70.5. None of these frozen models beats published sesc L=0.556. The models were not retuned. That is H3.",
    )
    _table(
        doc,
        ["System", "MAE", "Non-fl.", "L", "F2"],
        [
            ["Persist. LRP", "5.209", "4.287", "0.694", "0.739"],
            ["Unguarded XGB", "3.504", "9.333", "1.75e6", "0.000"],
            ["Residual XGB", "3.476", "9.289", "104", "0.017"],
            ["Floor hurdle", "3.177", "11.832", "70.5", "0.025"],
            ["Guarded ens.", "3.107", "6.750", "0.694", "0.739"],
        ],
    )
    _caption(doc, "Table III. Official test (n=2,167; 150 high-risk). L = MSE_HR / F2 after clipping below −6. Frozen before look.")
    _body(
        doc,
        "H4 is mixed. Combined log det of the snapshot covariance correlates with later |y − risk| (ρ=0.399). The max-risk gap is associated with ending at the floor (ρ=0.407; logistic test AUC 0.819) but with less later movement (ρ=−0.768). Quartile 1 has mean |y − risk|=13.03 and floor rate 0.56; quartiles 3–4 have floor rates above 0.95 (Figure 4).",
    )
    _figure(doc, *CAPTIONS[3])
    _body(
        doc,
        "The extra information in a T−48 CDM history is therefore early, and it is concentrated in later collapses to the dataset floor. Selecting on mean error ships a floor-calling policy. Selecting on ESA-style loss would have kept persistence, or the guard. Limitations: historical ESA-supported missions from 2015 to 2019, anonymized, not live ISRO. Manoeuvres are out of scope.",
    )

    _heading(doc, "Acknowledgements")
    _body(
        doc,
        "This work uses the ESA Space Debris Office Collision Avoidance Challenge dataset [3], [4]. ESA, ISRO, NASA, ISTRAC, and CCSDS did not endorse this project.",
        first=False,
    )
    _heading(doc, "References")
    refs = [
        "[1] Consultative Committee for Space Data Systems, Conjunction Data Message (CCSDS 508.0-B-1), 2013.",
        "[2] M. D. Hejduk and L. C. Johnson, “Satellite conjunction assessment risk analysis for dilution region events,” in Space Traffic Management Conference, 2016.",
        "[3] T. Uriot et al., “Spacecraft collision avoidance challenge: Design and results of a machine learning competition,” Astrodynamics, vol. 6, pp. 121–140, 2022.",
        "[4] ESA Advanced Concepts Team and Space Debris Office, Collision Avoidance Challenge dataset. Zenodo, 2021, doi: 10.5281/zenodo.4463683.",
        "[5] M. Duncan, J. Wysack, and J. Frisbee, “Collision probability forecasting using a Monte Carlo simulation,” in AMOS, 2014.",
        "[6] S. Alfano, “Relating position uncertainty to maximum conjunction probability,” AAS 03-548, 2003.",
        "[7] J. Lei et al., “Distribution-free predictive inference for regression,” J. Amer. Statist. Assoc., vol. 113, no. 523, pp. 1094–1111, 2018.",
        "[8] J. Stauch et al., “Contextual predictive model for early identification of high-covariance conjunctions,” J. Astronaut. Sci., vol. 73, p. 19, 2026.",
    ]
    for t in refs:
        p = _p(doc, t, size=8, space_after=2)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)

    doc.save(str(OUT_DOCX))
    print("Wrote", OUT_DOCX)


def _zip_dir(src: Path, dest: Path) -> None:
    dest.parent.mkdir(parents=True, exist_ok=True)
    if dest.exists():
        dest.unlink()
    with zipfile.ZipFile(dest, "w", zipfile.ZIP_DEFLATED) as zf:
        for p in sorted(src.rglob("*")):
            if p.is_file():
                zf.write(p, p.relative_to(src).as_posix())


def _run_skill(script: Path, *args: str) -> None:
    subprocess.check_call(
        [sys.executable, str(script), *args],
        cwd=str(script.parent),
    )


def fill_pptx() -> None:
    if UNPACK.exists():
        shutil.rmtree(UNPACK)
    UNPACK.mkdir(parents=True)
    with zipfile.ZipFile(PPTX_TMPL) as zf:
        zf.extractall(UNPACK)

    after = "slide2.xml"
    for _ in range(3):
        _run_skill(ADD_SLIDE, str(UNPACK), "slide2.xml", "--after", after)
        slides = sorted(
            UNPACK.glob("ppt/slides/slide*.xml"),
            key=lambda p: int(re.search(r"(\d+)", p.name).group(1)),
        )
        after = slides[-1].name

    pres_path = UNPACK / "ppt" / "presentation.xml"
    rels_path = UNPACK / "ppt" / "_rels" / "presentation.xml.rels"
    rels = rels_path.read_text(encoding="utf-8")
    m = re.search(
        r'<Relationship Id="(rId\d+)" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide" Target="slides/slide3.xml"/>',
        rels,
    )
    if not m:
        raise SystemExit("Could not find table slide (slide3) in template rels")
    table_rid = m.group(1)
    pres = pres_path.read_text(encoding="utf-8")
    pres = re.sub(rf'<p:sldId id="\d+" r:id="{table_rid}"/>', "", pres)
    pres_path.write_text(pres, encoding="utf-8")
    _run_skill(CLEAN_SLIDE, str(UNPACK))

    figure_slides = []
    for slide in sorted(
        (UNPACK / "ppt" / "slides").glob("slide*.xml"),
        key=lambda p: int(re.search(r"(\d+)", p.name).group(1)),
    ):
        xml = slide.read_text(encoding="utf-8")
        if "<p:pic>" in xml:
            figure_slides.append(slide)
    if len(figure_slides) != 4:
        raise SystemExit(f"Expected 4 figure slides, found {len(figure_slides)}")

    media = UNPACK / "ppt" / "media"
    media.mkdir(exist_ok=True)
    for i, (slide, (png_name, caption)) in enumerate(zip(figure_slides, CAPTIONS), start=1):
        dest = media / f"fig{i}.png"
        shutil.copy2(FIGS / png_name, dest)
        rel = UNPACK / "ppt" / "slides" / "_rels" / f"{slide.stem}.xml.rels"
        rel.write_text(
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideLayout" Target="../slideLayouts/slideLayout3.xml"/>'
            f'<Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="../media/fig{i}.png"/>'
            "</Relationships>",
            encoding="utf-8",
        )
        xml = slide.read_text(encoding="utf-8")
        xml = xml.replace(f"<a:t>{PLACEHOLDER_CAPTION}</a:t>", f"<a:t>{caption}</a:t>")
        slide.write_text(xml, encoding="utf-8")

    title_xml = (UNPACK / "ppt" / "slides" / "slide1.xml").read_text(encoding="utf-8")
    title_xml = title_xml.replace("CJSJ Figures Template", "CDM information decay")
    title_xml = title_xml.replace("2025-2026", "2026-2027")
    (UNPACK / "ppt" / "slides" / "slide1.xml").write_text(title_xml, encoding="utf-8")

    _run_skill(CLEAN_SLIDE, str(UNPACK))
    _zip_dir(UNPACK, OUT_PPTX)
    print("Wrote", OUT_PPTX)


if __name__ == "__main__":
    fill_docx()
    fill_pptx()
